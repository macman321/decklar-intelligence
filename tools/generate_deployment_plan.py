#!/usr/bin/env python3
"""
Decklar Deployment Plan Generator — Produces branded .docx documents.
Quality benchmark: examples/SE_Schneider_Deployment_Plan.docx

Usage:
    python3 generate_deployment_plan.py --session session.json --output plan.docx --kickoff-date "May 5, 2026"
    python3 generate_deployment_plan.py --answers answers.json --output plan.docx --kickoff-date "May 5, 2026"

Via import:
    from generate_deployment_plan import build_document
    build_document(answers, output_path, notes, kickoff_date)
"""

import json, sys, os, re, argparse
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import parse_xml

# ─── Decklar Brand Colors ─────────────────────────────────────────────
PURPLE = RGBColor(0x60, 0x38, 0xFB)
GREEN  = RGBColor(0x00, 0xCC, 0x55)
DARK   = RGBColor(0x1F, 0x29, 0x37)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
AMBER  = RGBColor(0xF5, 0xA6, 0x23)
RED    = RGBColor(0xDC, 0x26, 0x26)
BLUE   = RGBColor(0x3B, 0x82, 0xF6)
LIGHT_PURPLE_BG = "6038FB"

# ─── Helpers ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading_xml = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}" w:val="clear"/>'
    shading = parse_xml(shading_xml)
    cell._tc.get_or_add_tcPr().append(shading)

def _set_cell_border(cell, color_hex, size="12", pos="left"):
    """Add a left border accent to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders_xml = (
        f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:{pos} w:val="single" w:sz="{size}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcBorders = parse_xml(tcBorders_xml)
    tcPr.append(tcBorders)

def _style_text(run, font_name="Calibri", size=11, bold=False, color=DARK, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_heading(doc, text, level=1, color=PURPLE, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _style_text(run, size=size, bold=True, color=color)
    p.space_after = Pt(12)
    p.space_before = Pt(18)
    return p


def _add_paragraph(doc, text, bold=False, color=DARK, size=11, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_text(run, size=size, bold=bold, color=color, italic=italic)
    p.space_after = Pt(space_after)
    return p


def _add_bullet(doc, text, color=DARK, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    run = p.add_run(text)
    _style_text(run, color=color)
    p.space_after = Pt(4)
    return p


def _add_alert(doc, icon, text, level="critical"):
    """Add a color-coded alert paragraph (watch-out style)."""
    colors = {"critical": RED, "high": AMBER, "tip": BLUE}
    color = colors.get(level, GRAY)
    color_hex = {"critical": "DC2626", "high": "F5A623", "tip": "3B82F6"}.get(level, "6B7280")

    # Use a table for the left-border accent effect
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.25)

    cell = table.cell(0, 0)
    _set_cell_border(cell, color_hex, size="24", pos="left")
    _set_cell_bg(cell, "F8F7FF")  # very light purple bg

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{icon}  {text}")
    _style_text(run, size=11, bold=True, color=color)
    p.space_after = Pt(8)
    p.space_before = Pt(6)

    doc.add_paragraph()  # spacing


def _add_data_table(doc, headers, rows):
    """Add a styled data table with purple header, alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False

    # Set column widths roughly evenly
    total_width = Inches(6.25)
    col_width = int(total_width)  # docx expects integer emu
    for col in table.columns:
        col.width = col_width // len(headers)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_bg(cell, "6038FB")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        _style_text(run, size=10, bold=True, color=WHITE)
        p.space_after = Pt(4)
        p.space_before = Pt(4)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        bg = "FFFFFF" if row_idx % 2 == 0 else "F3F0FF"  # alternating white / very light purple
        for i, val in enumerate(row_data):
            cell = cells[i]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _style_text(run, size=10, color=DARK)
            p.space_after = Pt(3)
            p.space_before = Pt(3)

    doc.add_paragraph()  # spacing after table


# ─── Watch-Out Rule Engine ────────────────────────────────────────────

def generate_watchouts(answers):
    """Generate all watch-outs based on answer values."""
    watchouts = []

    # Always present
    watchouts.append(("🔴", "All devices must be added to account BEFORE shipping", "critical"))
    watchouts.append(("⚠️", "Monitor first 3 shipments closely after go-live", "high"))

    if answers.get("tripType") == "Round-Trip":
        watchouts.append(("🔴", "Battery life critical on multi-month round trips — confirm dwell time and disable unnecessary sensor interrupts", "critical"))

    if "BLE" in answers.get("connectivity", ""):
        watchouts.append(("⚠️", "BLE may be restricted at customer sites — verify with IT/security before deployment", "high"))

    if "⚠️ Not yet checked" in answers.get("deviceExpiry", ""):
        watchouts.append(("⚠️", "Bee Label expiry has not been verified — check before go-live", "high"))

    if answers.get("coldChain") == "Yes":
        watchouts.append(("ℹ️", "Configure temperature alert thresholds before go-live and validate with customer QA", "tip"))

    if answers.get("humidityMonitoring") == "Yes":
        watchouts.append(("ℹ️", "Set meaningful humidity alert thresholds — consider ambient vs. packaging conditions", "tip"))

    if "geofence" in answers.get("shipmentStart", "").lower() or "geofence" in answers.get("shipmentComplete", "").lower():
        watchouts.append(("ℹ️", "Validate geofence radius and trigger accuracy with facility coordinates before first shipment", "tip"))

    if "pharma" in answers.get("pharmaIndustry", "").lower():
        watchouts.append(("⚠️", "Pharma/regulated industry — ensure validation documentation and compliance audit trail are configured", "high"))

    if answers.get("oceanDuration", 0) > 30:
        watchouts.append(("🔴", f"Ocean transit ~{answers.get('oceanDuration')} days — verify battery life covers full journey + margin", "critical"))

    if answers.get("numLanes", 0) > 10:
        watchouts.append(("ℹ️", f"High lane count ({answers.get('numLanes')}) — consider phased rollout to manage complexity", "tip"))

    return watchouts


def generate_recommended_actions(answers, notes=""):
    """Generate customer-specific recommended actions."""
    actions = []
    customer = answers.get("customerName", "Customer")
    device = answers.get("deviceType", "TBD")
    mode = answers.get("shipmentMode", "TBD")

    # Action 1: Device provisioning
    if device == "One-Time Use":
        actions.append(f"Order and provision {answers.get('numLanes', 'TBD')} Bee Label batches for {customer} — confirm batch expiry covers go-live + 3 months")
    elif device == "Reusable — Reverse Logistics":
        actions.append(f"Configure reusable Bee return logistics for {customer} — establish origin charging dock and reverse logistics partner")
    else:
        actions.append(f"Confirm device type and provisioning plan with {customer} — device type is TBD")

    # Action 2: Account setup
    actions.append(f"Create {customer} account in Honeycomb and pre-provision all users ({answers.get('numUsers', 'TBD')} estimated)")

    # Action 3: Geofence / triggers
    if "geofence" in answers.get("shipmentStart", "").lower():
        actions.append(f"Validate geofence coordinates for all {answers.get('numLanes', 'TBD')} origin/destination facilities with {customer} facilities team")

    # Action 4: Alerts
    if answers.get("coldChain") == "Yes":
        actions.append(f"Configure temperature alert thresholds ({answers.get('tempThresholds', 'TBD')}) in Honeycomb and test with {customer} QA")

    # Action 5: BLE check
    if "BLE" in answers.get("connectivity", ""):
        actions.append(f"Confirm BLE policy with {customer} IT/security — get written approval if needed")

    # Action 6: Public tracking
    if "Yes" in answers.get("publicTracking", ""):
        actions.append(f"Enable public tracking links for {customer} and provide link-generation training to their logistics team")

    # Action 7: Open items from notes
    if notes:
        # Extract anything that sounds like a follow-up
        followups = re.findall(r'(?i)(follow[- ]?up|todo|action item|open item|need to)[\s:]*([^\n]+)', notes)
        for _, item in followups[:2]:
            actions.append(f"Follow-up: {item.strip()}")

    # Always include these
    actions.append(f"Schedule go-live validation call with {customer} within 2 weeks of first active shipment")
    actions.append(f"Send deployment plan .docx to {answers.get('escalationContact', 'primary contact').split('—')[0].strip()} for review and sign-off")

    return actions


def generate_open_items(answers):
    """Extract TBDs and create open items with proper ownership."""
    items = []
    customer = answers.get("customerName", "Customer")
    contact = answers.get("escalationContact", "")
    contact_name = contact.split("—")[0].strip() if "—" in contact else contact
    jeff = answers.get("accountManager", "Jeffrey Calabro")

    tbd_keys = [k for k, v in answers.items() if v == "TBD" or str(v).strip().upper() == "TBD" or str(v).strip() == ""]

    for key in tbd_keys:
        # Map key to human label
        label_map = {
            "successCriteria": "Success criteria",
            "activationLocation": "Activation location",
            "lightConfirmed": "Facility lighting confirmation",
            "deviceExpiry": "Device expiry verification",
            "accessControls": "Access control configuration",
            "numUsers": "Number of platform users",
            "escalationContact": "Primary escalation contact",
            "numLanes": "Total lane count",
            "laneDuration": "Average lane duration",
            "tempThresholds": "Temperature alert thresholds",
            "sensorInterrupts": "Sensor interrupt configuration",
        }
        label = label_map.get(key, key.replace("_", " ").title())

        # Determine owner
        if key in ["successCriteria", "activationLocation", "lightConfirmed", "numUsers", "accessControls", "escalationContact"]:
            owner = contact_name if contact_name else customer
        else:
            owner = jeff

        items.append({
            "item": f"Confirm {label} for {customer}",
            "owner": owner,
            "status": "Pending"
        })

    # Always add these
    if answers.get("deviceExpiry", "").startswith("⚠️"):
        items.append({
            "item": f"Verify all Bee Label expiry dates for {customer}",
            "owner": jeff,
            "status": "Pending"
        })

    if not any("light" in i["item"].lower() for i in items):
        items.append({
            "item": f"Confirm activation facility lighting (>40 foot candles) for {customer}",
            "owner": jeff,
            "status": "Pending"
        })

    return items


# ─── Main Document Builder ──────────────────────────────────────────

def build_document(answers, output_path, notes="", kickoff_date=None):
    """Build a complete branded deployment plan .docx."""

    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────
    sections = doc.sections[0]
    sections.top_margin = Inches(0.75)
    sections.bottom_margin = Inches(0.75)
    sections.left_margin = Inches(0.875)
    sections.right_margin = Inches(0.875)

    customer = answers.get("customerName", "Customer")
    am = answers.get("accountManager", "Jeffrey Calabro")
    today = datetime.now().strftime("%B %d, %Y")
    kd = kickoff_date or today

    # ── Cover Block ──────────────────────────────────────────────────
    # Full-width purple table as cover
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cover_table.autofit = False
    cover_table.allow_autofit = False
    cover_table.columns[0].width = Inches(6.25)

    cover_cell = cover_table.cell(0, 0)
    _set_cell_bg(cover_cell, "6038FB")
    cover_cell.paragraphs[0].clear()  # clear default paragraph

    # Customer name — large green text
    p = cover_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(customer)
    _style_text(run, size=28, bold=True, color=GREEN)
    p.space_after = Pt(6)

    # Account Manager
    p = cover_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Account Manager: {am}")
    _style_text(run, size=12, color=WHITE)
    p.space_after = Pt(2)

    # Kickoff Date
    p = cover_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Kickoff Date: {kd}")
    _style_text(run, size=12, color=WHITE)
    p.space_after = Pt(2)

    # Today
    p = cover_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Generated: {today}")
    _style_text(run, size=11, italic=True, color=WHITE)
    p.space_after = Pt(12)

    # Source
    p = cover_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Source: Decklar AI Onboarding Portal")
    _style_text(run, size=10, italic=True, color=WHITE)
    p.space_after = Pt(4)

    doc.add_paragraph()  # spacing after cover

    # ── Primary Contact ──────────────────────────────────────────────
    _add_heading(doc, "Primary Contact", level=1, color=PURPLE, size=14)
    contact = answers.get("escalationContact", "TBD")
    if contact and contact != "TBD":
        if "—" in contact:
            name, email = contact.split("—", 1)
            _add_paragraph(doc, f"{name.strip()}", bold=True)
            _add_paragraph(doc, f"Email: {email.strip()}")
        else:
            _add_paragraph(doc, contact)
    else:
        _add_paragraph(doc, "TBD — to be confirmed with customer", italic=True, color=GRAY)

    # ── Success Criteria ─────────────────────────────────────────────
    _add_heading(doc, "Success Criteria", level=1, color=PURPLE, size=14)
    sc = answers.get("successCriteria", "")
    if sc and sc != "TBD":
        _add_paragraph(doc, sc)
    else:
        _add_paragraph(doc, "TBD — to be defined with customer during follow-up", italic=True, color=GRAY)

    if notes:
        _add_paragraph(doc, f"Notes: {notes}", italic=True, color=GRAY, size=10)

    # ── Shipment Configuration ─────────────────────────────────────
    _add_heading(doc, "Shipment Configuration", level=1, color=PURPLE, size=14)

    headers = ["Parameter", "Value"]
    rows = [
        ["Shipment Mode", answers.get("shipmentMode", "TBD")],
        ["Number of Lanes", str(answers.get("numLanes", "TBD"))],
        ["Average Lane Duration", f"{answers.get('laneDuration', 'TBD')} days"],
        ["Trip Type", answers.get("tripType", "TBD")],
        ["Shipment Start Trigger", answers.get("shipmentStart", "TBD")],
        ["Shipment Complete Trigger", answers.get("shipmentComplete", "TBD")],
    ]

    # Conditional rows
    if "Ocean" in answers.get("shipmentMode", ""):
        rows.append(["Ocean Duration", f"{answers.get('oceanDuration', 'TBD')} days"])
        rows.append(["Container Tracking", answers.get("containerTracking", "TBD")])

    if "Air" in answers.get("shipmentMode", ""):
        rows.append(["Airlines", answers.get("airlines", "TBD")])
        rows.append(["AWB Tracking", answers.get("awbTracking", "TBD")])

    if answers.get("shipmentMode") == "Rail":
        rows.append(["Rail Dwell Time", answers.get("railDwell", "TBD")])

    _add_data_table(doc, headers, rows)

    # ── Sensor & Device Configuration ─────────────────────────────────
    _add_heading(doc, "Sensor & Device Configuration", level=1, color=PURPLE, size=14)

    headers = ["Parameter", "Value"]
    rows = [
        ["Device Type", answers.get("deviceType", "TBD")],
        ["Cold Chain", answers.get("coldChain", "TBD")],
        ["Temperature Thresholds", answers.get("tempThresholds", "N/A")],
        ["Shock Monitoring", answers.get("shockMonitoring", "TBD")],
        ["Humidity Monitoring", answers.get("humidityMonitoring", "TBD")],
        ["Connectivity", answers.get("connectivity", "TBD")],
        ["Sensor Interrupts", answers.get("sensorInterrupts", "TBD")],
    ]
    _add_data_table(doc, headers, rows)

    # ── Account Capabilities ─────────────────────────────────────────
    _add_heading(doc, "Account Capabilities", level=1, color=PURPLE, size=14)

    headers = ["Capability", "Enabled"]
    rows = [
        ["Public Tracking Links", answers.get("publicTracking", "TBD")],
        ["e-Proof of Delivery / Departure", answers.get("eProof", "TBD")],
        ["Waypoint Readiness Alerts", answers.get("waypointAlerts", "TBD")],
        ["Detention / Demurrage Tracking", answers.get("detentionDemurrage", "TBD")],
        ["Vehicle & Asset Tracking", answers.get("vehicleAssets", "TBD")],
        ["Pharma / Regulated Industry Mode", answers.get("pharmaIndustry", "TBD")],
        ["Access Controls", answers.get("accessControls", "TBD")],
        ["Number of Platform Users", str(answers.get("numUsers", "TBD"))],
    ]
    _add_data_table(doc, headers, rows)

    # ── Deployment Details ───────────────────────────────────────────
    _add_heading(doc, "Deployment Details", level=1, color=PURPLE, size=14)

    headers = ["Detail", "Value"]
    rows = [
        ["Device Activation Location", answers.get("activationLocation", "TBD")],
        ["Facility Lighting Confirmed", answers.get("lightConfirmed", "TBD")],
        ["Device Expiry Verified", answers.get("deviceExpiry", "TBD")],
    ]
    _add_data_table(doc, headers, rows)

    # ── Watch-Outs ───────────────────────────────────────────────────
    _add_heading(doc, "Watch-Outs", level=1, color=PURPLE, size=14)
    _add_paragraph(doc, "Critical considerations for this deployment:", italic=True, color=GRAY, size=10)

    watchouts = generate_watchouts(answers)
    for icon, text, level in watchouts:
        _add_alert(doc, icon, text, level)

    # ── Recommended Actions ──────────────────────────────────────────
    _add_heading(doc, "Recommended Actions", level=1, color=PURPLE, size=14)
    _add_paragraph(doc, "Next steps for Jeffrey Calabro and the Decklar team:", italic=True, color=GRAY, size=10)

    actions = generate_recommended_actions(answers, notes)
    for i, action in enumerate(actions, 1):
        _add_bullet(doc, f"{i}. {action}")

    # ── Open Items ───────────────────────────────────────────────────
    _add_heading(doc, "Open Items", level=1, color=PURPLE, size=14)
    _add_paragraph(doc, "Items requiring follow-up before or during deployment:", italic=True, color=GRAY, size=10)

    open_items = generate_open_items(answers)
    if open_items:
        headers = ["Item", "Owner", "Status"]
        rows = [[i["item"], i["owner"], i["status"]] for i in open_items]
        _add_data_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "No open items — all configuration confirmed.", italic=True, color=GREEN)

    # ── Footer ───────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by Decklar AI Intelligence System")
    _style_text(run, size=9, italic=True, color=GRAY)
    run = p.add_run(f"  •  {today}")
    _style_text(run, size=9, italic=True, color=GRAY)

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(output_path)
    return output_path


# ─── CLI Entry Point ────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate Decklar Deployment Plan .docx")
    parser.add_argument("--session", help="Path to portal session JSON file")
    parser.add_argument("--answers", help="Path to answers JSON file")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    parser.add_argument("--kickoff-date", default=None, help="Kickoff date string")
    parser.add_argument("--notes", default="", help="Additional notes")
    args = parser.parse_args()

    answers = {}
    notes = args.notes

    if args.session:
        with open(args.session) as f:
            session = json.load(f)
            answers = session.get("answers", {})
            notes = session.get("notes", notes)
    elif args.answers:
        with open(args.answers) as f:
            answers = json.load(f)
    else:
        print("Error: provide --session or --answers")
        sys.exit(1)

    output = build_document(answers, args.output, notes, args.kickoff_date)
    print(f"Deployment plan generated: {output}")


if __name__ == "__main__":
    main()
